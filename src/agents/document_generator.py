"""
Document Generator Agent - Compiles trip data into formatted documents
Runs as the final step in a sequential workflow to summarize all collected information
Generates downloadable .docx files
"""

import os
from datetime import datetime
from typing import Dict, Any, List
from pathlib import Path
from loguru import logger
from .base_agent import BaseAgent

# Try to import python-docx for .docx generation
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. .docx generation will be disabled.")

# Output directory for generated documents
OUTPUT_DIR = Path(__file__).parent.parent.parent / "outputs"
OUTPUT_DIR.mkdir(exist_ok=True)


class DocumentGeneratorAgent(BaseAgent):
    """
    Document Generator Agent for compiling trip planning documents.

    This agent runs at the end of the sequential workflow to compile
    all collected data (weather, visa, currency, flights, hotels, itinerary)
    into a well-formatted trip document.
    """

    def __init__(self):
        super().__init__(
            name="document_generator",
            description="Compiles trip data into formatted documents"
        )

        # Register A2A message handlers
        self.register_message_handler("compile_document", self._handle_compile_request)

    async def _execute_impl(self, input_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Generate a compiled trip document from all collected data.

        Args:
            input_data: Contains all trip data sections:
                - destination, origin, dates, travelers, budget, interests
                - weather_data, visa_data, currency_data
                - flight_data, hotel_data, itinerary_data

        Returns:
            Formatted trip document
        """
        # Extract trip overview
        destination = input_data.get("destination", "")
        origin = input_data.get("origin", "")
        start_date = input_data.get("start_date", "")
        end_date = input_data.get("end_date", "")
        travelers = input_data.get("travelers", 2)
        budget = input_data.get("budget", 0)
        interests = input_data.get("interests", "")

        # Calculate trip duration
        days = self._calculate_days(start_date, end_date)

        # Extract data sections from previous agents
        weather_data = input_data.get("weather_data", {})
        visa_data = input_data.get("visa_data", {})
        currency_data = input_data.get("currency_data", {})
        flight_data = input_data.get("flight_data", {})
        hotel_data = input_data.get("hotel_data", {})
        itinerary_data = input_data.get("itinerary_data", {})

        # Compile the document
        document = self._compile_document(
            destination=destination,
            origin=origin,
            start_date=start_date,
            end_date=end_date,
            days=days,
            travelers=travelers,
            budget=budget,
            interests=interests,
            weather_data=weather_data,
            visa_data=visa_data,
            currency_data=currency_data,
            flight_data=flight_data,
            hotel_data=hotel_data,
            itinerary_data=itinerary_data
        )

        # Generate .docx file if available
        docx_path = None
        docx_filename = None
        if DOCX_AVAILABLE:
            docx_result = self._generate_docx(document)
            docx_path = docx_result.get("file_path")
            docx_filename = docx_result.get("filename")
            document["docx_filename"] = docx_filename
            document["docx_path"] = docx_path
            logger.info(f"[DOCUMENT] Generated .docx file: {docx_filename}")

        # Send notification to orchestrator
        self.send_message(
            to_agent="orchestrator",
            message_type="document_generated",
            content={
                "destination": destination,
                "filename": document["filename"],
                "docx_filename": docx_filename
            }
        )

        return {
            "status": "success",
            "document": document,
            "docx_available": DOCX_AVAILABLE and docx_path is not None,
            "docx_filename": docx_filename,
            "docx_path": str(docx_path) if docx_path else None,
            "download_url": f"/download/{docx_filename}" if docx_filename else None,
            "metadata": {
                "generated_at": datetime.now().isoformat(),
                "sections_included": list(document["sections"].keys()),
                "formats": ["markdown", "docx"] if docx_filename else ["markdown"]
            }
        }

    def _calculate_days(self, start_date: str, end_date: str) -> int:
        """Calculate number of days between dates."""
        try:
            d1 = datetime.strptime(start_date, "%Y-%m-%d")
            d2 = datetime.strptime(end_date, "%Y-%m-%d")
            return (d2 - d1).days + 1
        except:
            return 7

    def _compile_document(
        self,
        destination: str,
        origin: str,
        start_date: str,
        end_date: str,
        days: int,
        travelers: int,
        budget: float,
        interests: str,
        weather_data: Dict[str, Any],
        visa_data: Dict[str, Any],
        currency_data: Dict[str, Any],
        flight_data: Dict[str, Any],
        hotel_data: Dict[str, Any],
        itinerary_data: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Compile all data into a structured document."""

        # Create filename-safe destination
        safe_destination = destination.lower().replace(' ', '_').replace(',', '')

        document = {
            "title": f"Trip to {destination}",
            "filename": f"{safe_destination}_trip_{start_date}.md",
            "format": "markdown",

            "trip_overview": {
                "destination": destination,
                "origin": origin,
                "dates": f"{start_date} to {end_date}",
                "duration": f"{days} days / {days - 1} nights",
                "travelers": travelers,
                "budget": f"${budget:,.2f}" if budget else "Not specified",
                "interests": interests
            },

            "sections": {
                "weather_and_packing": self._format_weather_section(weather_data),
                "visa_requirements": self._format_visa_section(visa_data),
                "currency_and_budget": self._format_currency_section(currency_data),
                "flight_options": self._format_flight_section(flight_data),
                "hotel_options": self._format_hotel_section(hotel_data),
                "itinerary": self._format_itinerary_section(itinerary_data, days, interests)
            },

            "summary": self._generate_summary(
                destination, days, travelers, budget,
                weather_data, visa_data, currency_data
            )
        }

        return document

    def _format_weather_section(self, weather_data: Dict[str, Any]) -> Dict[str, Any]:
        """Format weather data for document."""
        if not weather_data:
            return {"available": False, "note": "Weather data not collected"}

        return {
            "available": True,
            "current_conditions": weather_data.get("conditions", "N/A"),
            "temperature": weather_data.get("temperature", "N/A"),
            "feels_like": weather_data.get("feels_like", "N/A"),
            "humidity": weather_data.get("humidity", "N/A"),
            "packing_suggestions": weather_data.get("packing_suggestions", [])
        }

    def _format_visa_section(self, visa_data: Dict[str, Any]) -> Dict[str, Any]:
        """Format visa data for document."""
        if not visa_data:
            return {"available": False, "note": "Visa data not collected"}

        # Check if domestic travel
        if visa_data.get("travel_type") == "domestic":
            return {
                "available": True,
                "travel_type": "domestic",
                "visa_required": False,
                "message": visa_data.get("message", "Domestic travel - no visa needed")
            }

        return {
            "available": True,
            "travel_type": "international",
            "visa_required": visa_data.get("required", "Unknown"),
            "visa_type": visa_data.get("visa_type", "N/A"),
            "max_stay": visa_data.get("max_stay", "N/A"),
            "requirements": visa_data.get("requirements", []),
            "notes": visa_data.get("notes", "")
        }

    def _format_currency_section(self, currency_data: Dict[str, Any]) -> Dict[str, Any]:
        """Format currency and budget data for document."""
        if not currency_data:
            return {"available": False, "note": "Currency data not collected"}

        # Check if domestic travel
        if currency_data.get("travel_type") == "domestic":
            return {
                "available": True,
                "travel_type": "domestic",
                "currency_exchange_needed": False,
                "message": currency_data.get("message", "Same currency used")
            }

        # Extract currency info
        currency_info = currency_data.get("currency_info", currency_data)
        budget_breakdown = currency_data.get("budget_breakdown", {})

        return {
            "available": True,
            "from_currency": currency_info.get("from_currency", "USD"),
            "to_currency": currency_info.get("to_currency", "N/A"),
            "exchange_rate": currency_info.get("rate", "N/A"),
            "converted_amount": currency_info.get("converted", "N/A"),
            "budget_breakdown": budget_breakdown,
            "saving_tips": currency_data.get("saving_tips", []),
            "payment_recommendations": currency_data.get("payment_recommendations", {})
        }

    def _format_flight_section(self, flight_data: Dict[str, Any]) -> Dict[str, Any]:
        """Format flight data for document."""
        if not flight_data:
            return {"available": False, "note": "Flight data not collected"}

        return {
            "available": True,
            "search_params": flight_data.get("search_params", {}),
            "options": flight_data.get("options", []),
            "instruction": flight_data.get("instruction_for_llm", ""),
            "booking_tips": flight_data.get("booking_tips", [])
        }

    def _format_hotel_section(self, hotel_data: Dict[str, Any]) -> Dict[str, Any]:
        """Format hotel data for document."""
        if not hotel_data:
            return {"available": False, "note": "Hotel data not collected"}

        # Check if from Amadeus API
        if "hotels" in hotel_data:
            return {
                "available": True,
                "source": hotel_data.get("source", "amadeus_api"),
                "hotels": hotel_data.get("hotels", []),
                "booking_tips": hotel_data.get("booking_tips", [])
            }

        return {
            "available": True,
            "source": "llm_knowledge",
            "info": hotel_data,
            "booking_tips": hotel_data.get("booking_tips", [])
        }

    def _format_itinerary_section(
        self,
        itinerary_data: Dict[str, Any],
        days: int,
        interests: str
    ) -> Dict[str, Any]:
        """Format itinerary data for document."""
        if not itinerary_data:
            return {
                "available": False,
                "note": "Itinerary data not collected",
                "instruction": f"Generate a {days}-day itinerary focusing on: {interests}"
            }

        return {
            "available": True,
            "days": itinerary_data.get("days", days),
            "interests": itinerary_data.get("interests", interests),
            "instruction": itinerary_data.get("instruction", ""),
            "activities": itinerary_data.get("activities", [])
        }

    def _generate_summary(
        self,
        destination: str,
        days: int,
        travelers: int,
        budget: float,
        weather_data: Dict[str, Any],
        visa_data: Dict[str, Any],
        currency_data: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Generate trip summary with key highlights."""

        highlights = []
        reminders = []

        # Weather highlight
        if weather_data:
            temp = weather_data.get("temperature", "")
            conditions = weather_data.get("conditions", "")
            if temp and conditions:
                highlights.append(f"Expect {conditions} with temperatures around {temp}")

        # Visa reminder
        if visa_data:
            if visa_data.get("travel_type") == "international":
                if visa_data.get("required"):
                    reminders.append("Visa required - apply in advance")
                else:
                    highlights.append("No visa required for this trip")

        # Budget highlight
        if currency_data and currency_data.get("currency_info"):
            currency_info = currency_data.get("currency_info", {})
            rate = currency_info.get("rate")
            to_currency = currency_info.get("to_currency")
            if rate and to_currency:
                highlights.append(f"Current exchange rate: 1 USD = {rate} {to_currency}")

        # Standard reminders
        reminders.extend([
            "Confirm all bookings before departure",
            "Keep copies of important documents",
            "Check passport validity (6+ months recommended)",
            "Inform your bank of travel plans"
        ])

        return {
            "destination": destination,
            "duration": f"{days} days",
            "travelers": travelers,
            "budget": f"${budget:,.2f}" if budget else "Not specified",
            "highlights": highlights,
            "reminders": reminders
        }

    async def _handle_compile_request(self, message: Dict[str, Any]) -> Dict[str, Any]:
        """Handle A2A compile document request."""
        logger.info(f"[A2A] Received compile request from {message.get('from_agent')}")
        return await self.execute(message.get("content", {}))

    def _generate_docx(self, document: Dict[str, Any]) -> Dict[str, Any]:
        """
        Generate a professionally formatted .docx file from the trip document.

        Args:
            document: The compiled document dictionary

        Returns:
            Dict with file_path and filename
        """
        if not DOCX_AVAILABLE:
            return {"error": "python-docx not installed"}

        try:
            doc = Document()

            # ===== TITLE PAGE =====
            title = doc.add_heading(document.get("title", "Trip Plan"), 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Trip overview
            overview = document.get("trip_overview", {})
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(f"Dates: {overview.get('dates', 'N/A')}").bold = True
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(f"Travelers: {overview.get('travelers', 'N/A')} | Budget: {overview.get('budget', 'N/A')}")

            doc.add_page_break()

            # ===== TABLE OF CONTENTS (manual) =====
            doc.add_heading("Table of Contents", level=1)
            toc_items = [
                "1. Trip Overview",
                "2. Weather & Packing",
                "3. Visa Requirements",
                "4. Currency & Budget",
                "5. Flight Options",
                "6. Hotel Options",
                "7. Day-by-Day Itinerary",
                "8. Important Reminders"
            ]
            for item in toc_items:
                doc.add_paragraph(item)

            doc.add_page_break()

            # ===== SECTION 1: TRIP OVERVIEW =====
            doc.add_heading("1. Trip Overview", level=1)
            overview_table = doc.add_table(rows=6, cols=2)
            overview_table.style = 'Table Grid'
            overview_data = [
                ("Destination", overview.get("destination", "N/A")),
                ("Origin", overview.get("origin", "N/A")),
                ("Travel Dates", overview.get("dates", "N/A")),
                ("Duration", overview.get("duration", "N/A")),
                ("Travelers", str(overview.get("travelers", "N/A"))),
                ("Budget", overview.get("budget", "N/A"))
            ]
            for i, (label, value) in enumerate(overview_data):
                overview_table.rows[i].cells[0].text = label
                overview_table.rows[i].cells[1].text = str(value)

            doc.add_paragraph()

            # ===== SECTION 2: WEATHER & PACKING =====
            sections = document.get("sections", {})
            doc.add_heading("2. Weather & Packing", level=1)
            weather = sections.get("weather_and_packing", {})
            if weather.get("available"):
                doc.add_paragraph(f"Current Conditions: {weather.get('current_conditions', 'N/A')}")
                doc.add_paragraph(f"Temperature: {weather.get('temperature', 'N/A')}")
                doc.add_paragraph(f"Humidity: {weather.get('humidity', 'N/A')}")
                if weather.get("packing_suggestions"):
                    doc.add_heading("Packing Suggestions:", level=2)
                    for item in weather.get("packing_suggestions", []):
                        doc.add_paragraph(f"• {item}", style='List Bullet')
            else:
                doc.add_paragraph("Weather data not available.")

            # ===== SECTION 3: VISA REQUIREMENTS =====
            doc.add_heading("3. Visa Requirements", level=1)
            visa = sections.get("visa_requirements", {})
            if visa.get("available"):
                if visa.get("travel_type") == "domestic":
                    doc.add_paragraph("This is domestic travel - no visa required.")
                else:
                    doc.add_paragraph(f"Visa Required: {'Yes' if visa.get('visa_required') else 'No'}")
                    if visa.get("visa_type"):
                        doc.add_paragraph(f"Visa Type: {visa.get('visa_type')}")
                    if visa.get("max_stay"):
                        doc.add_paragraph(f"Maximum Stay: {visa.get('max_stay')}")
            else:
                doc.add_paragraph("Visa information not available.")

            # ===== SECTION 4: CURRENCY & BUDGET =====
            doc.add_heading("4. Currency & Budget", level=1)
            currency = sections.get("currency_and_budget", {})
            if currency.get("available"):
                if currency.get("travel_type") == "domestic":
                    doc.add_paragraph("Domestic travel - no currency exchange needed.")
                else:
                    doc.add_paragraph(f"From: {currency.get('from_currency', 'USD')}")
                    doc.add_paragraph(f"To: {currency.get('to_currency', 'N/A')}")
                    doc.add_paragraph(f"Exchange Rate: {currency.get('exchange_rate', 'N/A')}")

                    # Budget breakdown
                    breakdown = currency.get("budget_breakdown", {})
                    if breakdown:
                        doc.add_heading("Budget Breakdown:", level=2)
                        for category, amount in breakdown.items():
                            doc.add_paragraph(f"• {category}: ${amount}" if isinstance(amount, (int, float)) else f"• {category}: {amount}")
            else:
                doc.add_paragraph("Currency information not available.")

            # ===== SECTION 5: FLIGHT OPTIONS =====
            doc.add_heading("5. Flight Options", level=1)
            flights = sections.get("flight_options", {})
            if flights.get("available"):
                options = flights.get("options", [])
                for i, flight in enumerate(options[:3], 1):
                    doc.add_heading(f"Option {i}", level=2)
                    if isinstance(flight, dict):
                        doc.add_paragraph(f"Airline: {flight.get('airline', 'N/A')}")
                        doc.add_paragraph(f"Price: {flight.get('price', 'N/A')}")
                    else:
                        doc.add_paragraph(str(flight))
            else:
                doc.add_paragraph("Flight information not available.")

            # ===== SECTION 6: HOTEL OPTIONS =====
            doc.add_heading("6. Hotel Options", level=1)
            hotels = sections.get("hotel_options", {})
            if hotels.get("available"):
                hotel_list = hotels.get("hotels", [])
                for i, hotel in enumerate(hotel_list[:3], 1):
                    doc.add_heading(f"Option {i}", level=2)
                    if isinstance(hotel, dict):
                        doc.add_paragraph(f"Name: {hotel.get('name', 'N/A')}")
                        doc.add_paragraph(f"Price: {hotel.get('price', 'N/A')}")
                        doc.add_paragraph(f"Rating: {hotel.get('rating', 'N/A')}")
                    else:
                        doc.add_paragraph(str(hotel))
            else:
                doc.add_paragraph("Hotel information not available.")

            # ===== SECTION 7: ITINERARY =====
            doc.add_heading("7. Day-by-Day Itinerary", level=1)
            itinerary = sections.get("itinerary", {})
            if itinerary.get("available"):
                activities = itinerary.get("activities", [])
                if activities:
                    for activity in activities:
                        doc.add_paragraph(f"• {activity}", style='List Bullet')
                else:
                    doc.add_paragraph(itinerary.get("instruction", "Itinerary to be generated."))
            else:
                doc.add_paragraph(itinerary.get("instruction", "Itinerary not available."))

            # ===== SECTION 8: REMINDERS =====
            doc.add_heading("8. Important Reminders", level=1)
            summary = document.get("summary", {})
            for reminder in summary.get("reminders", []):
                doc.add_paragraph(f"✓ {reminder}", style='List Bullet')

            # ===== FOOTER =====
            doc.add_paragraph()
            doc.add_paragraph()
            footer = doc.add_paragraph()
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer.add_run(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')} by AI Vacation Planner").italic = True

            # Save the document
            safe_dest = document.get("trip_overview", {}).get("destination", "trip").lower().replace(" ", "_").replace(",", "")
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{safe_dest}_{timestamp}.docx"
            file_path = OUTPUT_DIR / filename

            doc.save(str(file_path))

            return {
                "filename": filename,
                "file_path": str(file_path),
                "success": True
            }

        except Exception as e:
            logger.error(f"[DOCUMENT] Failed to generate .docx: {e}")
            return {"error": str(e), "success": False}
